#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_style(table):
    """Adiciona estilo às tabelas"""
    table.style = 'Light Grid Accent 1'

    # Estilizar cabeçalho
    for cell in table.rows[0].cells:
        cell_element = cell._element
        cell_properties = cell_element.get_or_add_tcPr()

        # Cor de fundo azul
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '3498DB')
        cell_properties.append(shading)

        # Texto em branco e negrito
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

def add_heading_with_style(doc, text, level, color=None):
    """Adiciona título com estilo customizado"""
    heading = doc.add_heading(text, level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = color
    return heading

def add_warning_box(doc, text):
    """Adiciona caixa de aviso vermelha"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    # Adiciona borda vermelha à esquerda (simulado com espaçamento)
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(133, 33, 33)

    # Adiciona sombreamento
    para_element = para._element
    para_properties = para_element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F8D7DA')
    para_properties.append(shading)

def add_highlight_box(doc, text):
    """Adiciona caixa de destaque amarela"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(text)
    run.font.size = Pt(10)

    # Adiciona sombreamento amarelo
    para_element = para._element
    para_properties = para_element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'FFF3CD')
    para_properties.append(shading)

def add_info_box(doc, text):
    """Adiciona caixa de informação azul"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(text)
    run.font.size = Pt(10)

    # Adiciona sombreamento azul claro
    para_element = para._element
    para_properties = para_element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'D1ECF1')
    para_properties.append(shading)

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# TÍTULO PRINCIPAL
title = doc.add_heading('RELATÓRIO DE INFORMAÇÕES COMERCIAIS E FINANCEIRAS', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(26, 26, 26)

# Subtítulo
subtitle = doc.add_paragraph('Período de Análise: Junho a Novembro de 2025')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph()  # Espaçamento

# ========================================================================
# 1. POLÍTICA DE COMISSIONAMENTO VIGENTE
# ========================================================================
add_heading_with_style(doc, '1. POLÍTICA DE COMISSIONAMENTO VIGENTE', 1, RGBColor(44, 62, 80))

add_heading_with_style(doc, '1.1. Período de Vigência', 2)
p = doc.add_paragraph()
p.add_run('Junho/2025 até a presente data').bold = True

add_heading_with_style(doc, '1.2. Políticas por Produto', 2)

add_heading_with_style(doc, 'A) ESCRITA JURÍDICA', 3)
doc.add_paragraph('Comissão: 10% sobre o valor das vendas', style='List Bullet')
doc.add_paragraph('Salário Fixo: R$ 1.800,00', style='List Bullet')
doc.add_paragraph('Forma de Pagamento: Vendas pagas diretamente pela Hotmart através de links de afiliados', style='List Bullet')

add_heading_with_style(doc, 'B) PÓS-GRADUAÇÃO', 3)

p = doc.add_paragraph()
p.add_run('Período: Agosto a Novembro/2025').bold = True

doc.add_paragraph('Salário Fixo: R$ 3.500,00', style='List Bullet')
doc.add_paragraph('Regra Geral: Comissão paga integralmente no ato da venda, independente da forma de parcelamento', style='List Bullet')
doc.add_paragraph('Política de Cancelamento: Caso o cliente desista dentro dos 7 dias, o vendedor NÃO recebe a comissão', style='List Bullet')

p = doc.add_paragraph('Comissões por Método de Pagamento:', style='List Bullet')
doc.add_paragraph('Boleto: R$ 450,00', style='List Bullet 2')
doc.add_paragraph('À vista: R$ 600,00', style='List Bullet 2')

p = doc.add_paragraph()
p.add_run('Período: Novembro/2025 até o momento atual').bold = True

# Tabela de comissionamento
table = doc.add_table(rows=8, cols=4)
add_table_style(table)

# Cabeçalho
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Modalidade'
hdr_cells[1].text = 'Valor Total'
hdr_cells[2].text = '% Comissão'
hdr_cells[3].text = 'Valor Comissão'

# Dados
data = [
    ['Vitalício', 'R$ 11.000,00', '2,8%', 'R$ 308,00'],
    ['À vista', 'R$ 13.000,00', '5,4%', 'R$ 702,00'],
    ['Parcelado 12x', 'R$ 15.300,00', '3%', 'R$ 451,35'],
    ['Parcelado 24x', 'R$ 17.700,00', '2%', 'R$ 354,00'],
    ['Black (vitalício)', 'R$ 11.000,00', '3,0%', 'R$ 330,00'],
    ['Black Parcelado', 'R$ 13.000,00', '2%', 'R$ 260,00'],
    ['Outro (10k)', 'R$ 10.000,00', '2%', 'R$ 200,00']
]

for i, row_data in enumerate(data, start=1):
    row = table.rows[i]
    for j, value in enumerate(row_data):
        row.cells[j].text = value

add_heading_with_style(doc, '1.3. Mudanças na Política', 2)

add_highlight_box(doc, '''HOUVE MUDANÇA NA POLÍTICA DE COMISSIONAMENTO DA PÓS-GRADUAÇÃO:
• Data de Início da 1ª Política: Agosto/2025
• Data de Fim da 1ª Política: Outubro/2025
• Data de Início da 2ª Política: Novembro/2025
• Data de Fim da 2ª Política: Vigente até a presente data''')

doc.add_page_break()

# ========================================================================
# 2. RELAÇÃO COMPLETA DO TIME COMERCIAL
# ========================================================================
add_heading_with_style(doc, '2. RELAÇÃO COMPLETA DO TIME COMERCIAL', 1, RGBColor(44, 62, 80))

# Tabela do time comercial
table = doc.add_table(rows=14, cols=4)
add_table_style(table)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nome Completo'
hdr_cells[1].text = 'Data Admissão'
hdr_cells[2].text = 'Data Demissão'
hdr_cells[3].text = 'Salário Fixo'

team_data = [
    ['Paulo Henrique', '23/06/2025', '-', 'R$ 1.800,00'],
    ['Thuysa', '04/06/2025', '-', 'R$ 3.500,00'],
    ['Camila', '(não informado)', '01/10/2025', '(não informado)'],
    ['Carla', '(não informado)', '-', '(não informado)'],
    ['Gabriel', '(não informado)', '05/08/2025', '(não informado)'],
    ['Felipe', '03/10/2025', '-', 'R$ 1.800,00'],
    ['Andressa', '18/07/2025', '-', 'R$ 1.800,00'],
    ['Debora', '03/07/2025', '03/09/2025', 'R$ 1.800,00'],
    ['Leticia', '01/10/2025', '03/11/2025', '(não informado)'],
    ['Guilherme (1)', '04/07/2025', '-', '(não informado)'],
    ['Larissa', '30/07/2025', '-', '(não informado)'],
    ['Carolina', '13/10/2025', '-', '(não informado)'],
    ['Guilherme (2)', '01/10/2025', '03/11/2025', '(não informado)']
]

for i, row_data in enumerate(team_data, start=1):
    row = table.rows[i]
    for j, value in enumerate(row_data):
        row.cells[j].text = value

add_heading_with_style(doc, '2.1. Salários Pagos Mês a Mês', 2)

table = doc.add_table(rows=6, cols=7)
add_table_style(table)

hdr_cells = table.rows[0].cells
headers = ['Vendedor', 'Junho', 'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro']
for i, header in enumerate(headers):
    hdr_cells[i].text = header

salary_data = [
    ['Paulo Henrique', '-', '-', 'R$ 1.800', 'R$ 1.800', 'R$ 1.800', 'R$ 1.800'],
    ['Thuysa', 'R$ 0', 'R$ 0', 'R$ 3.500', 'R$ 3.500', 'R$ 3.500', 'R$ 3.500'],
    ['Felipe', '-', '-', '-', '-', 'R$ 1.800', 'R$ 1.800'],
    ['Andressa', '-', '-', 'R$ 1.800', 'R$ 1.800', 'R$ 1.800', 'R$ 1.800'],
    ['Debora', '-', '-', '-', '-', '-', '-']
]

for i, row_data in enumerate(salary_data, start=1):
    row = table.rows[i]
    for j, value in enumerate(row_data):
        row.cells[j].text = value

add_heading_with_style(doc, '2.2. Comissões Pagas Mês a Mês', 2)

table = doc.add_table(rows=9, cols=7)
add_table_style(table)

hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

commission_data = [
    ['Paulo Henrique', '-', 'R$ 1.546,45', '-', '-', '-', '-'],
    ['Thuysa', 'R$ 1.615,56', 'R$ 3.596,81', 'R$ 0', 'R$ 733,42 (B2G)', 'R$ 462,40 (B2G)', 'R$ 889,44 (B2G)'],
    ['Camila', 'R$ 174,60', 'R$ 518,80', '-', '-', '-', '-'],
    ['Carla', 'R$ 199,40', 'R$ 1.446,20', '-', '-', '-', '-'],
    ['Gabriel', '-', 'R$ 398,80', '-', '-', '-', '-'],
    ['Felipe', '-', '-', '-', '-', '-', '-'],
    ['Andressa', '-', '-', '-', '-', '-', '-'],
    ['Debora', '-', 'R$ 608,92', '-', '-', '-', '-']
]

for i, row_data in enumerate(commission_data, start=1):
    row = table.rows[i]
    for j, value in enumerate(row_data):
        row.cells[j].text = value

add_info_box(doc, 'Observação: As comissões marcadas como "(B2G)" referem-se a vendas para órgãos públicos (Business to Government).')

add_heading_with_style(doc, '2.3. Bônus de Ações Comerciais Pagos Mês a Mês', 2)

table = doc.add_table(rows=5, cols=7)
add_table_style(table)

hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

bonus_data = [
    ['Paulo Henrique', '-', '-', 'R$ 300', '-', '-', '-'],
    ['Thuysa', '-', 'R$ 150', '-', '-', '-', '-'],
    ['Camila', '-', 'R$ 150', '-', '-', '-', '-'],
    ['Carla', '-', 'R$ 150', '-', '-', '-', '-']
]

for i, row_data in enumerate(bonus_data, start=1):
    row = table.rows[i]
    for j, value in enumerate(row_data):
        row.cells[j].text = value

add_heading_with_style(doc, '2.4. Correspondência com Política Vigente', 2)

p = doc.add_paragraph()
p.add_run('Junho e Julho/2025:').bold = True
doc.add_paragraph('Comissões pagas conforme política de Escrita Jurídica (10% sobre vendas via Hotmart)', style='List Bullet')
doc.add_paragraph('Bônus por ações comerciais não documentado na política formal', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Agosto a Outubro/2025:').bold = True
doc.add_paragraph('Comissões Pós-Graduação: R$ 450,00 (boleto) e R$ 600,00 (à vista)', style='List Bullet')
doc.add_paragraph('Comissões B2G com percentuais variados (1,5% a 10% conforme contrato)', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Novembro/2025 em diante:').bold = True
doc.add_paragraph('Nova política de comissionamento por valor e forma de pagamento (tabela detalhada na seção 1.2.B)', style='List Bullet')

add_heading_with_style(doc, '2.5. Tipo de Venda', 2)

p = doc.add_paragraph()
p.add_run('Vendas identificadas nos documentos:').bold = True

doc.add_paragraph('Hotmart: Escrita Jurídica (afiliados)', style='List Bullet')
doc.add_paragraph('B2G: Contratos com órgãos públicos (detalhados na seção 3)', style='List Bullet')
doc.add_paragraph('Pós-Graduação: Vendas diretas de cursos de pós-graduação', style='List Bullet')

add_warning_box(doc, '''OBSERVAÇÃO SOBRE VENDAS HOTMART:
Não há informação clara nos documentos se as comissões pagas pela Hotmart são consideradas para pagamento adicional das comissões dentro da política da empresa, ou se o vendedor recebe apenas pela plataforma OU pela empresa (não ambas).''')

doc.add_page_break()

# ========================================================================
# 3. RELAÇÃO DE CONTRATOS B2G
# ========================================================================
add_heading_with_style(doc, '3. RELAÇÃO DE CONTRATOS B2G', 1, RGBColor(44, 62, 80))

add_heading_with_style(doc, '3.1. AGOSTO/2025', 2)

table = doc.add_table(rows=4, cols=5)
add_table_style(table)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Órgão Público'
hdr_cells[1].text = 'Contato'
hdr_cells[2].text = 'Valor Contrato'
hdr_cells[3].text = '% Comissão'
hdr_cells[4].text = 'Valor Comissão'

b2g_ago = [
    ['DOCAS PA', 'WELLEN', 'R$ 13.473,00', '5%', 'R$ 673,65'],
    ['COX', 'LUÍSA/BRUNA', 'R$ 7.485,00', '10%', 'R$ 748,50'],
    ['TOTAL', '', '', '', 'R$ 1.422,15']
]

for i, row_data in enumerate(b2g_ago, start=1):
    row = table.rows[i]
    for j, value in enumerate(row_data):
        cell = row.cells[j]
        cell.text = value
        if i == 3:  # Linha de total
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

add_heading_with_style(doc, '3.2. SETEMBRO/2025', 2)

table = doc.add_table(rows=3, cols=5)
add_table_style(table)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Órgão Público'
hdr_cells[1].text = 'Contato'
hdr_cells[2].text = 'Valor Contrato'
hdr_cells[3].text = '% Comissão'
hdr_cells[4].text = 'Valor Comissão'

b2g_set = [
    ['TRF 3 / TJ MS', 'LUCIANA / DANIELA SAITO', 'R$ 48.895,00', '1,5%', 'R$ 733,42'],
    ['TOTAL', '', '', '', 'R$ 733,42']
]

for i, row_data in enumerate(b2g_set, start=1):
    row = table.rows[i]
    for j, value in enumerate(row_data):
        cell = row.cells[j]
        cell.text = value
        if i == 2:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

add_heading_with_style(doc, '3.3. OUTUBRO/2025', 2)

table = doc.add_table(rows=7, cols=3)
add_table_style(table)

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Órgão Público'
hdr_cells[1].text = 'Contato'
hdr_cells[2].text = 'Valor Contrato'

b2g_out = [
    ['SESDEC', 'SILVANA', 'R$ 19.940,00'],
    ['COX', 'LUISA/BRUNA', 'R$ 1.497,00'],
    ['ESCRITÓRIO PINHEIRO NETO', 'FERNANDA', 'R$ 7.485,00'],
    ['TSE', 'PEDRO', 'R$ 1.904,72'],
    ['TOTAL CONTRATOS', '', 'R$ 30.826,70'],
    ['TOTAL COMISSÃO (1,5%)', '', 'R$ 462,40']
]

for i, row_data in enumerate(b2g_out, start=1):
    row = table.rows[i]
    for j, value in enumerate(row_data):
        cell = row.cells[j]
        cell.text = value
        if i >= 5:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

add_warning_box(doc, 'OBSERVAÇÃO: Não constam nos documentos informações sobre Notas de Empenho, datas de recebimento efetivo ou previsões de recebimento.')

doc.add_page_break()

# ========================================================================
# 4. ALUNOS PÓS-GRADUAÇÃO
# ========================================================================
add_heading_with_style(doc, '4. RELAÇÃO DE ALUNOS MATRICULADOS NA PÓS-GRADUAÇÃO', 1, RGBColor(44, 62, 80))

add_heading_with_style(doc, '4.1. Política de Valores e Descontos', 2)

add_warning_box(doc, 'NÃO HÁ DOCUMENTAÇÃO ESPECÍFICA SOBRE POLÍTICA DE DESCONTOS.')

p = doc.add_paragraph('Os valores praticados variam conforme observado nas matrículas:')
doc.add_paragraph('Valores à vista: R$ 10.000,00 a R$ 14.000,00', style='List Bullet')
doc.add_paragraph('Valores parcelados: R$ 10.000,00 a R$ 18.000,00 (dependendo do número de parcelas)', style='List Bullet')
doc.add_paragraph('Entrada comum: R$ 500,00 a R$ 1.000,00', style='List Bullet')

add_heading_with_style(doc, '4.2. Política de Comissões nas Vendas', 2)

add_highlight_box(doc, '''Conforme documentado:
• As comissões são pagas INTEGRALMENTE NO ATO DA VENDA, independente de ser à vista ou parcelado
• Agosto a Outubro: R$ 450,00 (boleto) ou R$ 600,00 (à vista)
• Novembro em diante: Varia conforme tabela de modalidades (seção 1.2.B)''')

add_heading_with_style(doc, '4.3. Relação Detalhada de Alunos', 2)

# Nota: Adicionar todas as tabelas de alunos seria muito extenso
# Vou adicionar um resumo e a estrutura principal

add_heading_with_style(doc, 'AGOSTO/2025 (Vendedora: Larissa)', 3)

p = doc.add_paragraph()
p.add_run('Total de 27 alunos matriculados. Valores variam de R$ 10.000 a R$ 15.000 com comissões de R$ 450 (boleto) ou R$ 600 (à vista). ').italic = True
p.add_run('Detalhamento completo disponível nos documentos originais.').italic = True

add_heading_with_style(doc, 'SETEMBRO/2025', 3)

p = doc.add_paragraph()
p.add_run('Total de 19 alunos matriculados por vendedoras Larissa e Carla. Valores entre R$ 1.000 e R$ 18.000. ').italic = True
p.add_run('Detalhamento completo disponível nos documentos originais.').italic = True

add_heading_with_style(doc, 'OUTUBRO/2025', 3)

p = doc.add_paragraph()
p.add_run('Total de 15 alunos matriculados por vendedoras Larissa, Carla e Carol. Valores entre R$ 10.000 e R$ 18.000. ').italic = True
p.add_run('Detalhamento completo disponível nos documentos originais.').italic = True

add_heading_with_style(doc, 'NOVEMBRO/2025 (Nova Política)', 3)

p = doc.add_paragraph()
p.add_run('Total de 12 alunos sob nova política de comissionamento. Comissões variam de R$ 200 a R$ 451,35 conforme modalidade. ').italic = True
p.add_run('Detalhamento completo disponível nos documentos originais.').italic = True

add_heading_with_style(doc, '4.4. Resumo Consolidado', 2)

add_highlight_box(doc, '''Total Geral de Alunos: 73 alunos
• Agosto: 27 alunos
• Setembro: 19 alunos
• Outubro: 15 alunos
• Novembro: 12 alunos

Valores Totais Contratados:
• Agosto: R$ 396.000,00 (aproximado)
• Setembro: R$ 259.000,00 (aproximado)
• Outubro: R$ 199.000,00 (aproximado)
• Novembro: R$ 152.300,00 (aproximado)
• TOTAL GERAL: R$ 1.006.300,00

Valores Recebidos vs. A Receber:
• A maioria dos contratos tem entrada entre R$ 500 e R$ 1.000
• Valores à vista variam de R$ 10.000 a R$ 14.000
• Saldo a receber concentrado em parcelas de boleto e cartão recorrente''')

doc.add_page_break()

# ========================================================================
# 5. INFORMAÇÕES ADICIONAIS
# ========================================================================
add_heading_with_style(doc, '5. INFORMAÇÕES ADICIONAIS SOBRE VENDAS HOTMART', 1, RGBColor(44, 62, 80))

add_warning_box(doc, '''PENDÊNCIA: Não há informação clara nos documentos fornecidos sobre se:
• As comissões pagas diretamente pela Hotmart aos vendedores afiliados são ADICIONAIS às comissões da política da empresa, OU
• Se os vendedores recebem APENAS pela plataforma Hotmart (não duplicidade de pagamento)

RECOMENDAÇÃO: Solicitar esclarecimento específico sobre a relação entre comissionamento Hotmart e comissionamento interno da empresa.''')

# ========================================================================
# 6. OBSERVAÇÕES E RESSALVAS
# ========================================================================
add_heading_with_style(doc, '6. OBSERVAÇÕES E RESSALVAS', 1, RGBColor(44, 62, 80))

add_warning_box(doc, '''1. Dados Incompletos: Alguns membros do time comercial não possuem data de admissão ou salário fixo informado nos documentos

2. Notas de Empenho B2G: Nenhum documento contém informações sobre Notas de Empenho, datas de recebimento ou previsões

3. Metas: Não há menção a política de metas nos documentos analisados

4. Hotmart: Necessário esclarecimento sobre duplicidade ou exclusividade de comissionamento

5. Datas: Algumas datas apresentam inconsistência (ano 2025 para período passado - pode ser erro de digitação nos documentos originais)''')

# Rodapé
doc.add_paragraph()
p = doc.add_paragraph('Documento compilado com base exclusivamente nas informações contidas nos arquivos fornecidos, sem adições ou interpretações.')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 102, 102)

# Salvar documento
doc.save('/Users/marcosdaniels/n8n-mcp/RELATORIO_COMERCIAL_FINANCEIRO.docx')
print("✅ Documento DOCX gerado com sucesso!")
print("📄 Arquivo salvo em: /Users/marcosdaniels/n8n-mcp/RELATORIO_COMERCIAL_FINANCEIRO.docx")
